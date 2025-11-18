"""基于功能解读要点汇总，调用大模型生成连贯完整的专题报告，并导出 DOCX。"""
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import json
import asyncio
import time
import traceback

from ..utils.setting.paths import bucket, ensure_bucket
from ..utils.logging.logging import (
    setup_logger,
    log_module_start,
    log_success,
    log_error,
    log_save_success,
)
from ..utils.setting.env_loader import get_api_key


def _get_date_folder(start_date: str, end_date: Optional[str]) -> str:
    """
    获取日期文件夹名，兼容单日与范围
    
    Args:
        start_date (str): 开始日期
        end_date (Optional[str]): 结束日期
    
    Returns:
        str: 形如 YYYY-MM-DD 或 YYYY-MM-DD_YYYY-MM-DD
    """
    return f"{start_date}_{end_date}" if end_date else start_date


def _ordered_functions() -> List[Tuple[str, str]]:
    """
    返回功能的读取顺序及中文名称
    
    Returns:
        List[Tuple[str, str]]: [(功能文件夹英文, 中文显示名)]
    """
    return [
        ("volume", "声量"),
        ("trends", "发布趋势"),
        ("classification", "分类"),
        ("attitude", "态度"),
        ("publishers", "发布者"),
        ("geography", "地域"),
        ("keywords", "关键词"),
        ("contentanalyze", "内容分析"),
        ("topic", "主题分析"),
    ]


def _ordered_channels() -> List[str]:
    """
    返回渠道读取顺序
    
    Returns:
        List[str]: 渠道列表
    """
    return ["总体"]


def _read_explain_text(explain_file: Path, logger) -> Optional[str]:
    """
    读取解读JSON中的文字部分
    
    Args:
        explain_file (Path): 解读文件路径
    
    Returns:
        Optional[str]: 文本，读取失败或字段缺失返回None
    """
    try:
        if not explain_file.exists():
            return None
        with open(explain_file, "r", encoding="utf-8") as f:
            data = json.load(f)
        text = data.get("explain")
        if isinstance(text, str) and text.strip():
            return text.strip()
        return None
    except Exception:
        log_error(logger, f"读取解读失败: {explain_file} | {traceback.format_exc()}", "Report")
        return None


def _collect_sections(topic: str, date_folder: str, logger) -> List[Tuple[str, str]]:
    """收集各功能【总体】下 *_rag_enhanced.json 的文字要点。

    Args:
        topic: 专题
        date_folder: 日期文件夹
        logger: 日志器

    Returns:
        [(功能中文, 文本)] 列表
    """
    explain_root = bucket("explain", topic, date_folder)
    sections: List[Tuple[str, str]] = []
    if not explain_root.exists():
        return sections
    for func_en, func_cn in _ordered_functions():
        func_dir = explain_root / func_en / "总体"
        explain_file = func_dir / f"{func_en}_rag_enhanced.json"
        text = _read_explain_text(explain_file, logger)
        if text:
            sections.append((func_cn, text))
    return sections


def _sections_to_block(sections: List[Tuple[str, str]]) -> str:
    """将功能要点拼接为单一文本块，保留功能名标签。"""
    parts: List[str] = []
    for func_cn, text in sections:
        parts.append(f"【{func_cn}】\n{text}".strip())
    return "\n\n".join(parts)


def _load_prompt_yaml(topic: str, logger) -> Dict:
    """读取专题提示词 YAML；若不存在则返回内置默认模板。"""
    try:
        import yaml  # type: ignore
    except Exception:
        log_error(logger, "缺少依赖 PyYAML，将使用内置模板", "Report")
        return {
            "system": (
                "你是一名资深中文编辑与舆情分析报告撰写者。你的任务是对已整理好的多段‘原文要点’做轻度编辑："
                "在不改变事实与重点的前提下，仅增加必要的衔接与过渡，使全文连贯、易读。严禁大幅压缩、改写或重新解读内容。"
            ),
            "user_template": (
                "请基于如下专题 {topic} 在 {date_folder} 的多功能原文要点，输出一份通顺连贯的中文报告。\n"
                "要求：\n"
                "- 保持各部分的原始信息与表达顺序，不重排、不删减关键信息。\n"
                "- 仅加入必要的衔接/过渡语与起承转合，减少文字的阻塞感。\n"
                "- 保留各部分的功能说明（如【地域】、【声量】等），可做轻微格式优化。\n"
                "- 输出可读的 Markdown（使用二、三级标题），不含图片/表格/代码。\n\n"
                "【原文要点（按顺序给出）】\n{sections_text}"
            ),
        }

    project_root = Path(__file__).resolve().parents[2]
    yaml_path = project_root / "configs" / "prompt" / "report" / "专题" / f"{topic}.yaml"
    if not yaml_path.exists():
        log_error(logger, f"未找到专题提示词: {yaml_path}，使用内置模板", "Report")
        return {
            "system": (
                "你是一名资深中文编辑与舆情分析报告撰写者。你的任务是对已整理好的多段‘原文要点’做轻度编辑："
                "在不改变事实与重点的前提下，仅增加必要的衔接与过渡，使全文连贯、易读。严禁大幅压缩、改写或重新解读内容。"
            ),
            "user_template": (
                "请基于如下专题 {topic} 在 {date_folder} 的多功能原文要点，输出一份通顺连贯的中文报告。\n"
                "要求：\n"
                "- 保持各部分的原始信息与表达顺序，不重排、不删减关键信息。\n"
                "- 仅加入必要的衔接/过渡语与起承转合，减少文字的阻塞感。\n"
                "- 保留各部分的功能说明（如【地域】、【声量】等），可做轻微格式优化。\n"
                "- 输出可读的 Markdown（二、三级标题），不含图片/表格/代码。\n\n"
                "【原文要点（按顺序给出）】\n{sections_text}"
            ),
        }
    try:
        import yaml  # type: ignore
        with open(yaml_path, "r", encoding="utf-8") as f:
            data = yaml.safe_load(f)  # type: ignore
        if not isinstance(data, dict):
            raise ValueError("提示词 YAML 结构需为对象字典")
        return data
    except Exception:
        log_error(logger, f"读取提示词失败: {yaml_path} | {traceback.format_exc()}", "Report")
        return {
            "system": (
                "你是一名资深中文编辑与舆情分析报告撰写者。你的任务是对已整理好的多段‘原文要点’做轻度编辑："
                "在不改变事实与重点的前提下，仅增加必要的衔接与过渡，使全文连贯、易读。严禁大幅压缩、改写或重新解读内容。"
            ),
            "user_template": (
                "请基于如下专题 {topic} 在 {date_folder} 的多功能原文要点，输出一份通顺连贯的中文报告。\n"
                "要求：\n"
                "- 保持各部分的原始信息与表达顺序，不重排、不删减关键信息。\n"
                "- 仅加入必要的衔接/过渡语与起承转合，减少文字的阻塞感。\n"
                "- 保留各部分的功能说明（如【地域】、【声量】等），可做轻微格式优化。\n"
                "- 输出可读的 Markdown（二、三级标题），不含图片/表格/代码。\n\n"
                "【原文要点（按顺序给出）】\n{sections_text}"
            ),
        }


def _compose_llm_input(topic: str, date_folder: str, sections_text: str, tmpl: Dict) -> Dict:
    """基于模板拼装对话消息。"""
    system_text = str(tmpl.get("system", ""))
    user_template = str(tmpl.get("user_template", ""))
    user_text = user_template.format(topic=topic, date_folder=date_folder, sections_text=sections_text)
    return {
        "messages": [
            {"role": "system", "content": system_text},
            {"role": "user", "content": user_text},
        ]
    }


async def _llm_call_report(messages: Dict, logger, model: Optional[str] = None, timeout: float = 60.0, max_retries: int = 2) -> Optional[str]:
    """通过 DashScope 兼容聊天接口调用大模型，返回文本。"""
    import aiohttp  # type: ignore
    import random

    api_key = get_api_key()
    if not api_key:
        log_error(logger, "未获取到 API Key，请配置环境变量", "Report")
        return None

    url = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
    hdrs = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": model or "qwen-turbo", **messages, "temperature": 0.3}

    # 记录提示长度，便于诊断是否因体量过大导致超时
    try:
        total_chars = 0
        for m in messages.get("messages", []):
            if isinstance(m, dict):
                total_chars += len(str(m.get("content", "")))
        log_success(logger, f"报告提示长度：{total_chars} chars", "Report")
    except Exception:
        pass

    attempt = 0
    while attempt <= max_retries:
        attempt += 1
        try:
            t0 = time.perf_counter()
            # 针对长文本逐次放宽超时：60s → 96s → 154s（默认），上限为 3x
            eff_timeout = min(timeout * (1.6 ** (attempt - 1)), timeout * 3)
            client_timeout = aiohttp.ClientTimeout(total=eff_timeout, connect=30, sock_connect=30, sock_read=eff_timeout)
            async with aiohttp.ClientSession(timeout=client_timeout) as sess:
                async with sess.post(url, headers=hdrs, json=payload) as resp:
                    text = await resp.text()
                    if resp.status != 200:
                        log_error(logger, f"大模型HTTP错误 {resp.status}: {text}", "Report")
                        raise RuntimeError(f"HTTP {resp.status}")
                    data = json.loads(text)
                    content = (
                        data.get("choices", [{}])[0]
                        .get("message", {})
                        .get("content")
                    )
                    if isinstance(content, str) and content.strip():
                        dt = (time.perf_counter() - t0) * 1000
                        log_success(logger, f"LLM生成成功，用时 {dt:.0f} ms", "Report")
                        return content.strip()
                    log_error(logger, "大模型返回为空或格式异常", "Report")
                    return None
        except Exception:
            if attempt <= max_retries:
                backoff = min(1.5 ** attempt, 10.0) + random.uniform(0.0, 0.6)
                log_error(logger, f"LLM调用失败，重试 {attempt}/{max_retries}（超时={eff_timeout:.0f}s）: {traceback.format_exc()}", "Report")
                await asyncio.sleep(backoff)
            else:
                log_error(logger, f"LLM调用失败且达最大重试: {traceback.format_exc()}", "Report")
                return None


def _add_inline_runs(p, text: str) -> None:
    """
    在段落中按 Markdown 基础内联格式写入文本：支持 **加粗**、*斜体*、`行内代码`。

    Args:
        p: python-docx 段落对象
        text (str): 原始文本
    """
    import re
    from docx.shared import Pt
    pattern = re.compile(r"(\*\*.+?\*\*|\*[^*].+?\*|`.+?`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            run.font.name = "宋体"
            run.font.size = Pt(10.5)
        token = m.group(0)
        content = token
        run = p.add_run()
        if token.startswith("**") and token.endswith("**"):
            content = token[2:-2]
            run.bold = True
            run.text = content
            run.font.name = "宋体"
            run.font.size = Pt(10.5)
        elif token.startswith("*") and token.endswith("*"):
            content = token[1:-1]
            run.italic = True
            run.text = content
            run.font.name = "宋体"
            run.font.size = Pt(10.5)
        elif token.startswith("`") and token.endswith("`"):
            content = token[1:-1]
            run.text = content
            run.font.name = "Consolas"
            run.font.size = Pt(10.5)
        pos = m.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        run.font.name = "宋体"
        run.font.size = Pt(10.5)


def _render_markdown_to_doc(doc, md_text: str) -> None:
    """
    将 Markdown 文本渲染到 Word 文档：支持 ##/### 标题、无序/有序列表、引用、内联加粗/斜体/行内代码。

    Args:
        doc: python-docx 文档
        md_text (str): Markdown 文本
    """
    import re
    from docx.shared import Pt
    in_code_block = False
    for raw in md_text.splitlines():
        line = raw.rstrip("\n\r")
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            doc.add_paragraph("")
            continue
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.name = "宋体"
            run.font.size = Pt(14)
            continue
        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.name = "宋体"
            run.font.size = Pt(12)
            continue

        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, text)
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, text)
            continue

        if re.match(r"^\s*>\s+", line):
            text = re.sub(r"^\s*>\s+", "", line)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.name = "宋体"
            run.font.size = Pt(10.5)
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, line)

def _build_doc_from_text(topic: str, date_folder: str, body_text: str, output_path: Path, logger) -> bool:
    """将完整文本写入 Word 文档（简单段落排版）。"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn

        doc = Document()

        # 全局正文字体：宋体 5号（约10.5pt）
        normal_style = doc.styles["Normal"].font
        normal_style.name = "宋体"
        normal_style.size = Pt(10.5)
        try:
            normal_style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        except Exception:
            pass

        # 大标题
        title_para = doc.add_paragraph()
        run = title_para.add_run(f"{topic}/{date_folder}分析报告")
        run.bold = True
        run.font.name = "宋体"
        run.font.size = Pt(14)
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        except Exception:
            pass

        doc.add_paragraph("")

        # 渲染 Markdown 主体
        _render_markdown_to_doc(doc, body_text)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return True
    except ImportError as e:
        log_error(logger, f"缺少依赖 python-docx，请先安装: {e}", "Report")
        return False
    except Exception:
        log_error(logger, f"生成DOCX失败: {traceback.format_exc()}", "Report")
        return False


def _build_doc(topic: str, date_folder: str, sections: List[Tuple[str, str, str]], output_path: Path, logger) -> bool:
    """
    构建并保存DOCX文档
    
    Args:
        topic (str): 专题
        date_folder (str): 日期文件夹名
        sections (List[Tuple[str, str, str]]): (功能中文, 渠道中文, 正文) 列表
        output_path (Path): 输出文件路径
        logger: 日志器
    
    Returns:
        bool: 保存是否成功
    """
    try:
        # 延迟导入以避免无依赖时报错影响模块加载
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn

        doc = Document()

        # 全局正文字体：宋体 5号（约10.5pt）
        normal_style = doc.styles["Normal"].font
        normal_style.name = "宋体"
        normal_style.size = Pt(10.5)
        try:
            normal_style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        except Exception:
            pass

        # 大标题：宋体 4号（约14pt） 加粗
        title_para = doc.add_paragraph()
        run = title_para.add_run(f"{topic}/{date_folder}分析")
        run.bold = True
        run.font.name = "宋体"
        run.font.size = Pt(14)
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        except Exception:
            pass

        # 空行
        doc.add_paragraph("")

        # 各章节
        for func_cn, channel_cn, body in sections:
            # 子标题：宋体 4号 加粗
            h_para = doc.add_paragraph()
            h_run = h_para.add_run(f"{func_cn} - {channel_cn}")
            h_run.bold = True
            h_run.font.name = "宋体"
            h_run.font.size = Pt(14)
            try:
                h_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass

            # 正文：宋体 5号
            p = doc.add_paragraph(body)
            for r in p.runs:
                r.font.name = "宋体"
                r.font.size = Pt(10.5)
                try:
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                except Exception:
                    pass

            # 段落后空行
            doc.add_paragraph("")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return True
    except ImportError as e:
        log_error(logger, f"缺少依赖 python-docx，请先安装: {e}", "Report")
        return False
    except Exception as e:
        log_error(logger, f"生成DOCX失败: {e}", "Report")
        return False


def run_report(topic: str, start_date: str, end_date: Optional[str] = None) -> bool:
    """
    生成专题报告（DOCX）
    
    Args:
        topic (str): 专题名称
        start_date (str): 开始日期 (YYYY-MM-DD)
        end_date (Optional[str]): 结束日期 (YYYY-MM-DD)
    
    Returns:
        bool: 是否成功
    """
    date_folder = _get_date_folder(start_date, end_date)
    logger = setup_logger(f"Report_{topic}", date_folder)

    try:
        log_module_start(logger, "Report")

        # 1) 收集 *_rag_enhanced.json（仅“总体”）
        raw_sections = _collect_sections(topic, date_folder, logger)
        if not raw_sections:
            log_error(logger, "未找到任何可用的解读结果（*_rag_enhanced.json），报告未生成", "Report")
            return False

        sections_text = _sections_to_block(raw_sections)

        # 2) 读取提示词 YAML 并组装消息
        tmpl = _load_prompt_yaml(topic, logger)
        messages = _compose_llm_input(topic, date_folder, sections_text, tmpl)

        # 3) 读取 LLM 配置（若存在）
        project_root = Path(__file__).resolve().parents[2]
        llm_cfg_path = project_root / "configs" / "llm.yaml"
        model_name: Optional[str] = None
        timeout_s: float = 200.0  # 报告文本较长，默认更宽松
        max_retries: int = 2
        try:
            import yaml  # type: ignore
            if llm_cfg_path.exists():
                with open(llm_cfg_path, "r", encoding="utf-8") as f:
                    llm_cfg = yaml.safe_load(f)  # type: ignore
                if isinstance(llm_cfg, dict):
                    # 模型名优先级：report.model > report_model > model/chat_model
                    report_cfg = llm_cfg.get("report") if isinstance(llm_cfg.get("report"), dict) else None
                    model_name = (
                        (report_cfg or {}).get("model")
                        or llm_cfg.get("report_model")
                        or llm_cfg.get("model")
                        or llm_cfg.get("chat_model")
                    )
                    # 超时优先级：report.timeout > report_timeout > timeout > 默认
                    if report_cfg and "timeout" in report_cfg:
                        timeout_s = float(report_cfg.get("timeout", timeout_s))
                    else:
                        timeout_s = float(llm_cfg.get("report_timeout", llm_cfg.get("timeout", timeout_s)))
                    # 重试优先级：report.retries > report_retries > retries > 默认
                    if report_cfg and "retries" in report_cfg:
                        max_retries = int(report_cfg.get("retries", max_retries))
                    else:
                        max_retries = int(llm_cfg.get("report_retries", llm_cfg.get("retries", max_retries)))
            # 保护性下限：避免被设置得过小导致大文本易超时
            timeout_s = max(timeout_s, 90.0)
            log_success(logger, f"LLM配置：timeout={timeout_s:.0f}s, retries={max_retries}", "Report")
        except Exception:
            log_error(logger, f"读取 llm.yaml 失败，使用默认配置 | {traceback.format_exc()}", "Report")

        # 4) 调用大模型
        full_text: Optional[str] = asyncio.run(
            _llm_call_report(messages, logger, model=model_name, timeout=timeout_s, max_retries=max_retries)
        )

        # 5) 若 LLM 失败，回退为拼接文本直接输出
        if not full_text:
            log_error(logger, "LLM 生成失败，将回退为直接拼接文本", "Report")
            full_text = sections_text

        # 6) 写入 Word
        output_dir = ensure_bucket("reports", topic, date_folder)
        output_path = output_dir / f"{topic}_{date_folder}_报告.docx"
        ok = _build_doc_from_text(topic, date_folder, full_text, output_path, logger)
        if ok:
            log_success(logger, "报告生成完成（由 LLM 生产全文）", "Report")
        return ok

    except Exception:
        log_error(logger, f"报告生成失败: {traceback.format_exc()}", "Report")
        return False


